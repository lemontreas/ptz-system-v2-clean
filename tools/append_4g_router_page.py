from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("硬件安装说明_工作副本.docx")
OUTPUT = Path("硬件安装与4G路由器使用说明.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=10.5, bold=False, color=None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_compact_paragraph(container, text="", *, bold=False, size=10.5, before=0, after=2):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold)
    return paragraph


document = Document(SOURCE)

section = document.add_section(WD_SECTION.NEW_PAGE)
section.top_margin = Cm(1.25)
section.bottom_margin = Cm(1.25)
section.left_margin = Cm(1.55)
section.right_margin = Cm(1.55)
section.header_distance = Cm(0.6)
section.footer_distance = Cm(0.6)

title = document.add_paragraph()
title.paragraph_format.space_after = Pt(7)
title_run = title.add_run("8. 4G路由器安装与网络连接")
style_run(title_run, size=16, bold=True, color=(31, 78, 121))

intro = add_compact_paragraph(
    document,
    "4G路由器用于为现场设备提供局域网和互联网连接。首次使用前，请先安装4G卡，再连接路由器热点。",
    size=10.5,
    after=5,
)

layout = document.add_table(rows=1, cols=2)
layout.autofit = False
layout.columns[0].width = Cm(8.4)
layout.columns[1].width = Cm(8.4)
left, right = layout.rows[0].cells
left.width = Cm(8.4)
right.width = Cm(8.4)
for cell in (left, right):
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

add_compact_paragraph(left, "8.1 安装4G卡", bold=True, size=11.5, after=4)
photo = add_compact_paragraph(
    left,
    "【照片待补】\n请在此处放入“4G卡插槽位置及插卡方向”照片，并用红框或箭头标出卡槽。",
    size=10,
    before=2,
    after=5,
)
photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_shading(left, "F2F2F2")

steps = (
    "① 关闭路由器电源后再插卡。\n"
    "② 按卡槽标识确认芯片面和缺口方向，请勿强行插入。\n"
    "③ 插到底并确认固定后再通电。\n"
    "④ 等待4G/网络指示灯稳定；无法联网时，先检查卡是否装反、欠费或需要设置APN。"
)
add_compact_paragraph(left, steps, size=9.8, after=0)

add_compact_paragraph(right, "8.2 随设备提供的网络信息", bold=True, size=11.5, after=4)
info = right.add_table(rows=5, cols=2)
info.autofit = False
labels = ("热点名称（WiFi）", "热点密码", "默认网段", "管理员页面", "管理员账号/密码")
values = (
    "【待填写】",
    "【待填写】",
    "【待填写，如 192.168.8.0/24】",
    "连接热点后按下方步骤查看",
    "【待填写或查看机身铭牌】",
)
for row, label, value in zip(info.rows, labels, values):
    row.cells[0].width = Cm(3.1)
    row.cells[1].width = Cm(5.0)
    for cell in row.cells:
        set_cell_margins(cell, top=65, start=80, bottom=65, end=80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(row.cells[0], "D9EAF7")
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    style_run(p0.add_run(label), size=9, bold=True)
    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    style_run(p1.add_run(value), size=9)

add_compact_paragraph(
    right,
    "提示：热点名称、密码和管理信息通常印在路由器底部铭牌上；如已由设备方修改，以随设备提供的信息为准。",
    size=9.2,
    before=4,
    after=0,
)

add_compact_paragraph(document, "8.3 查看管理员页面及设备IP", bold=True, size=11.5, before=6, after=3)
guide = document.add_table(rows=4, cols=2)
guide.autofit = False
guide.columns[0].width = Cm(2.3)
guide.columns[1].width = Cm(14.7)
guide_rows = (
    ("连接热点", "使用电脑或手机连接上表中的路由器热点。"),
    (
        "查看管理地址",
        "Windows按 Win+R，输入 cmd 后运行 ipconfig；找到“无线局域网适配器 Wi-Fi”中的“默认网关”。默认网关通常就是路由器管理地址，例如 192.168.8.1。",
    ),
    (
        "打开管理员页面",
        "在浏览器地址栏输入 http://默认网关地址，使用机身铭牌或上表提供的管理员账号和密码登录。",
    ),
    (
        "查找设备IP",
        "进入“终端管理、已连接设备、DHCP客户端”或类似页面，按设备名称或MAC地址找到本设备；显示的IPv4地址就是当前设备IP。路由器重启后IP可能变化，需要固定时可在管理页面设置DHCP地址保留。",
    ),
)
for row, (label, text) in zip(guide.rows, guide_rows):
    row.cells[0].width = Cm(2.3)
    row.cells[1].width = Cm(14.7)
    for cell in row.cells:
        set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(row.cells[0], "D9EAF7")
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    style_run(p0.add_run(label), size=9.2, bold=True)
    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    style_run(p1.add_run(text), size=9.2)

warning = add_compact_paragraph(
    document,
    "注意：不要把“默认网关（路由器地址）”当成“设备IP”。访问设备页面或进行现场调试时，应使用管理员页面中分配给本设备的IPv4地址。",
    bold=True,
    size=9.5,
    before=5,
    after=0,
)
warning.paragraph_format.keep_together = True

document.save(OUTPUT)
print(OUTPUT.resolve())
